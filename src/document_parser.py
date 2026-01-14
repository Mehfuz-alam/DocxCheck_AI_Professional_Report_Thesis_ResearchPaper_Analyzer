# src/document_parser.py
from PyPDF2 import PdfReader
import docx
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import os

def parse_document(file_path):
    """
    Extract text and styles from PDF or DOCX.
    """
    text = ""
    styles = []

    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
            font_name = para.style.font.name if para.style.font else None
            font_size = para.style.font.size.pt if para.style.font and para.style.font.size else None
            styles.append({
                "text": para.text,
                "font_name": font_name,
                "font_size": font_size,
                "alignment": para.alignment
            })
    else:
        raise ValueError("Unsupported file type")

    return text, styles


def highlight_errors_docx(input_path, output_path, format_errors, structure_errors):
    """
    Highlight formatting and structure errors in DOCX only.
    PDF files are skipped (cannot modify directly).
    """
    if not input_path.endswith(".docx"):
        # Cannot highlight PDF directly
        return None

    doc = docx.Document(input_path)

    # Highlight formatting & structure errors
    for para in doc.paragraphs:
        snippet = para.text[:30]
        for err in format_errors + structure_errors:
            if snippet in err:
                if para.runs:
                    para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Add header/footer
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.text = "DocxCheck AI - Professional Report Analyzer"
        header_paragraph.alignment = 1  # center

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.text = "Page "
        footer_paragraph.alignment = 1

    doc.save(output_path)
    return output_path
